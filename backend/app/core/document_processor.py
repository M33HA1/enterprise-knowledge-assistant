"""
Document processing: parsing and chunking.

Supports PDF, DOCX, and TXT files.
Uses LangChain's RecursiveCharacterTextSplitter for intelligent chunking.

Design decisions:
  - RecursiveCharacterTextSplitter over simple splitting because it respects
    paragraph/sentence boundaries, producing more coherent chunks.
  - 500 char chunks with 50 char overlap balances retrieval precision with context.
  - Each chunk carries metadata (source file, page, chunk index, department)
    which enables RBAC filtering and source citation.
"""

import os
import logging
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A chunk of text with its metadata for embedding and retrieval."""
    content: str
    metadata: dict = field(default_factory=dict)
    # metadata keys: source, page, chunk_index, department, doc_id, file_type


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""
    filename: str
    content: str
    pages: List[str]
    file_type: str
    total_pages: int


class DocumentParser:
    """Parses PDF, DOCX, and TXT files into raw text."""

    @staticmethod
    def parse(file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            return DocumentParser._parse_pdf(path)
        elif ext == ".docx":
            return DocumentParser._parse_docx(path)
        elif ext == ".txt":
            return DocumentParser._parse_txt(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")

    @staticmethod
    def _parse_pdf(path: Path) -> ParsedDocument:
        """Extract text from PDF, preserving page boundaries."""
        import pymupdf  # PyMuPDF

        pages = []
        with pymupdf.open(str(path)) as doc:
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    pages.append(text.strip())

        if not pages:
            raise ValueError(f"No extractable text in PDF: {path.name}")

        return ParsedDocument(
            filename=path.name,
            content="\n\n".join(pages),
            pages=pages,
            file_type="pdf",
            total_pages=len(pages),
        )

    @staticmethod
    def _parse_docx(path: Path) -> ParsedDocument:
        """Extract text from DOCX."""
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        if not content.strip():
            raise ValueError(f"No extractable text in DOCX: {path.name}")

        return ParsedDocument(
            filename=path.name,
            content=content,
            pages=[content],  # DOCX doesn't have pages
            file_type="docx",
            total_pages=1,
        )

    @staticmethod
    def _parse_txt(path: Path) -> ParsedDocument:
        """Read plain text file."""
        content = path.read_text(encoding="utf-8")
        if not content.strip():
            raise ValueError(f"Empty text file: {path.name}")

        return ParsedDocument(
            filename=path.name,
            content=content,
            pages=[content],
            file_type="txt",
            total_pages=1,
        )


class DocumentChunker:
    """
    Splits parsed documents into chunks suitable for embedding.

    Uses RecursiveCharacterTextSplitter which tries to split at:
    1. Paragraph breaks (\\n\\n)
    2. Line breaks (\\n)
    3. Sentences (. ! ?)
    4. Words (spaces)
    5. Characters (last resort)
    """

    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

    def chunk_document(
        self,
        parsed_doc: ParsedDocument,
        department: str = "general",
        doc_id: Optional[str] = None,
    ) -> List[DocumentChunk]:
        """
        Split a parsed document into chunks with metadata.

        For PDFs, we chunk per-page to preserve page number metadata.
        For other formats, we chunk the entire content.
        """
        chunks = []
        chunk_index = 0

        if parsed_doc.file_type == "pdf":
            # Chunk each page separately to preserve page numbers
            for page_num, page_text in enumerate(parsed_doc.pages, start=1):
                page_chunks = self._splitter.split_text(page_text)
                for text in page_chunks:
                    chunks.append(DocumentChunk(
                        content=text,
                        metadata={
                            "source": parsed_doc.filename,
                            "page": page_num,
                            "chunk_index": chunk_index,
                            "department": department,
                            "doc_id": doc_id or parsed_doc.filename,
                            "file_type": parsed_doc.file_type,
                            "total_pages": parsed_doc.total_pages,
                        },
                    ))
                    chunk_index += 1
        else:
            # Chunk entire content for non-PDF
            text_chunks = self._splitter.split_text(parsed_doc.content)
            for text in text_chunks:
                chunks.append(DocumentChunk(
                    content=text,
                    metadata={
                        "source": parsed_doc.filename,
                        "page": 1,
                        "chunk_index": chunk_index,
                        "department": department,
                        "doc_id": doc_id or parsed_doc.filename,
                        "file_type": parsed_doc.file_type,
                        "total_pages": parsed_doc.total_pages,
                    },
                ))
                chunk_index += 1

        logger.info(
            f"Chunked '{parsed_doc.filename}': {len(chunks)} chunks "
            f"(size={self.chunk_size}, overlap={self.chunk_overlap})"
        )
        return chunks


def process_document(
    file_path: str,
    department: str = "general",
    doc_id: Optional[str] = None,
) -> List[DocumentChunk]:
    """
    Convenience function: parse + chunk a document in one call.

    Args:
        file_path: Path to the document file
        department: Department tag for RBAC filtering
        doc_id: Optional unique document identifier

    Returns:
        List of DocumentChunk ready for embedding
    """
    parser = DocumentParser()
    parsed = parser.parse(file_path)
    chunker = DocumentChunker()
    return chunker.chunk_document(parsed, department=department, doc_id=doc_id)
